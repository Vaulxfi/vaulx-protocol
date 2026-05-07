"""
Vaulx — 1-page company brief for Mercado Bitcoin meeting.
DOCX version (Word). Removes "Indicative Commercial Terms" (over-reach).

Sections:
  Header   · vaulx logo + tagline + meeting context
  1. The asymmetry we solve     · Brazilian credit ladder
  2. What is Vaulx              · protocol + traction
  3. How it works               · 4-step mechanism
  4. Unit economics             · borrower / capital providers / Vaulx revenue
  5. Where Mercado Bitcoin fits · 3 areas of cooperation
  6. What protects everyone     · 4 safeguards
  Footer  · team line · URLs · confidential

Output: Vaulx_MB_OnePager.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# === Brand colors =========================================================
INK         = RGBColor(0x1A, 0x1A, 0x1D)
INK_DARK    = RGBColor(0x0A, 0x0A, 0x0B)
ACCENT      = RGBColor(0x0E, 0x7C, 0x7B)
ACCENT_2    = RGBColor(0x2B, 0xA0, 0x9E)
MUTE        = RGBColor(0x6B, 0x6B, 0x70)
WARN        = RGBColor(0xB8, 0x41, 0x2C)
RULE        = RGBColor(0xD8, 0xD8, 0xDB)
PAPER_2     = RGBColor(0xF8, 0xF6, 0xF2)


# === Helpers ==============================================================
def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="D8D8DB", size=4):
    """Set individual cell borders. Each side: 'single', 'nil', or None (don't change)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is None:
            continue
        # remove existing
        existing = tc_borders.find(qn(f"w:{side}"))
        if existing is not None:
            tc_borders.remove(existing)
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tc_borders.append(b)


def set_cell_padding(cell, top=80, bottom=80, left=100, right=100):
    """Set cell padding in 1/20pt units (twips). 100 ≈ 5pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def add_horizontal_rule(paragraph, color="0E7C7B", size=12):
    """Add a horizontal line (bottom border) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def section_header(doc, num, title):
    """Add a numbered section header with teal accent + hairline below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    r = p.add_run(f"0{num}  ")
    r.font.name = "Helvetica"
    r.font.size = Pt(8)
    r.font.color.rgb = ACCENT
    r.font.bold = True

    r = p.add_run(title.upper())
    r.font.name = "Helvetica"
    r.font.size = Pt(10)
    r.font.color.rgb = INK_DARK
    r.font.bold = True

    add_horizontal_rule(p, color="D8D8DB", size=4)


def add_body(doc, text, size=9, color=INK, bold=False, italic=False, font="Helvetica",
             space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    p.alignment = alignment
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.italic = italic
    return p


# ============================================================================
# Build the document
# ============================================================================
doc = Document()

# Page setup — A4, narrow margins for one-page density
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.top_margin = Mm(14)
section.bottom_margin = Mm(14)
section.left_margin = Mm(15)
section.right_margin = Mm(15)


# ============================================================================
# HEADER  (logo left, meta right) — using a 2-column table
# ============================================================================
header_tbl = doc.add_table(rows=1, cols=2)
header_tbl.autofit = False
header_tbl.columns[0].width = Mm(110)
header_tbl.columns[1].width = Mm(70)

# LEFT — logo + tagline
left_cell = header_tbl.rows[0].cells[0]
left_cell.width = Mm(110)
set_cell_padding(left_cell, top=0, bottom=0, left=0, right=0)
set_cell_borders(left_cell, top="nil", bottom="nil", left="nil", right="nil")

# Logo image (use the prepared black-on-transparent version)
logo_path = "vaulx_logo_black.png"
logo_p = left_cell.paragraphs[0]
logo_p.paragraph_format.space_after = Pt(0)
if os.path.exists(logo_path):
    run = logo_p.add_run()
    run.add_picture(logo_path, height=Mm(10))

# Tagline below logo
tagline_p = left_cell.add_paragraph()
tagline_p.paragraph_format.space_before = Pt(2)
tagline_p.paragraph_format.space_after = Pt(0)
r = tagline_p.add_run("On-chain credit, secured by physical luxury collateral.")
r.font.name = "Helvetica"
r.font.size = Pt(8.5)
r.font.color.rgb = MUTE

# RIGHT — meta block (right-aligned)
right_cell = header_tbl.rows[0].cells[1]
right_cell.width = Mm(70)
set_cell_padding(right_cell, top=0, bottom=0, left=0, right=0)
set_cell_borders(right_cell, top="nil", bottom="nil", left="nil", right="nil")

for line, weight, size, color in [
    ("Prepared for: Mercado Bitcoin", True,  8.5, INK_DARK),
    ("Subject: Asset Tokenization & On-Chain Credit", False, 8.5, INK),
    ("Date: May 7, 2026", False, 8.5, MUTE),
    ("vaulx.fi · github.com/Vaulxfi", False, 8.0, ACCENT),
]:
    p = right_cell.paragraphs[0] if line.startswith("Prepared for") else right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(line)
    r.font.name = "Helvetica"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = weight

# Horizontal rule under header
ruler = doc.add_paragraph()
ruler.paragraph_format.space_before = Pt(2)
ruler.paragraph_format.space_after = Pt(0)
add_horizontal_rule(ruler, color="0A0A0B", size=10)


# ============================================================================
# 1. THE ASYMMETRY
# ============================================================================
section_header(doc, 1, "The asymmetry we solve")

# Asset-rich. Credit-trapped. (large serif headline)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Asset-rich. Credit-trapped.")
r.font.name = "Times New Roman"
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = INK_DARK

add_body(doc,
    "A São Paulo Rolex owner with a US$ 14,000 watch faces three credit options — all bad. "
    "Caixa Federal pawn lends 20% LTV at ~30% APR (scrap-metal valuation). A consumer loan starts "
    "at ~60% APR. A revolving credit-card balance hits ~400% APR. Meanwhile, on-chain USDC "
    "liquidity sits idle at 8% APR with no trustable rail to physical collateral.",
    size=9, space_before=0, space_after=4,
)

# Rate strip — 3-cell table
rate_tbl = doc.add_table(rows=1, cols=3)
rate_tbl.autofit = False
for col in rate_tbl.columns:
    col.width = Mm(60)
rates = [
    ("Caixa pawn (20% LTV)",    "~30% APR"),
    ("Consumer loan",           "~60% APR"),
    ("Credit-card revolving",   "~400% APR"),
]
for i, (label, val) in enumerate(rates):
    cell = rate_tbl.rows[0].cells[i]
    set_cell_bg(cell, "F8F6F2")
    set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")
    set_cell_padding(cell, top=80, bottom=80, left=140, right=140)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(label.upper())
    r.font.name = "Helvetica"
    r.font.size = Pt(7)
    r.font.color.rgb = MUTE
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(val)
    r.font.name = "Helvetica"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = WARN


# ============================================================================
# 2. WHAT IS VAULX
# ============================================================================
section_header(doc, 2, "What is Vaulx")

add_body(doc,
    "Vaulx is an on-chain credit protocol on Solana that connects asset-rich individuals in "
    "emerging markets with global on-chain liquidity — secured by physical luxury collateral "
    "(luxury watches first, then jewelry, gold, art) with deterministic on-chain liquidation.",
    size=9.5, space_before=2, space_after=3,
)

add_body(doc,
    "We do not take custody. We do not hold capital. We orchestrate licensed counterparties — "
    "independent custodians, certified appraisers, Lloyd's-class insurance — and route the loan "
    "through Solana smart contracts that release USDC only after the custodian signs custody-"
    "confirmation atomically on-chain.",
    size=9.5, space_before=0, space_after=4,
)

# Traction strip — 4 cell table
trac_tbl = doc.add_table(rows=1, cols=4)
trac_tbl.autofit = False
stats = [
    ("4",        "Anchor programs · live on Devnet"),
    ("45+",      "tests · all green · CI gating"),
    ("vaulx.fi", "frontend live · admin cockpit"),
    ("Q3 '26",   "mainnet target · 50 borrowers"),
]
for i, (big, sub) in enumerate(stats):
    cell = trac_tbl.rows[0].cells[i]
    set_cell_bg(cell, "F8F6F2")
    set_cell_borders(cell, top="nil", bottom="nil", left="single", right="nil",
                      color="0E7C7B", size=12)
    set_cell_padding(cell, top=70, bottom=70, left=100, right=80)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(big)
    r.font.name = "Helvetica"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = INK_DARK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(sub)
    r.font.name = "Helvetica"
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUTE


# ============================================================================
# 3. HOW IT WORKS
# ============================================================================
section_header(doc, 3, "How it works")

steps_tbl = doc.add_table(rows=1, cols=5)
steps_tbl.autofit = False
steps = [
    ("01", "Appraise",        "Trilateral appraisal — online API, qualified offline appraiser, and risk-officer review — sets the fair collateral value."),
    ("02", "Custody",         "Asset physically locked in licensed independent custody. Custodian signs custody-confirmation on-chain."),
    ("03", "cNFT mint",       "Custody confirmation triggers cNFT mint on Solana — embeds appraisal, custody, insurance, and redemption rights."),
    ("04", "Borrow",          "Borrower locks cNFT in the Vaulx contract; USDC drawn from Solana LPs and disbursed atomically. BRL via Pix off-ramp."),
    ("05", "Repay / Default", "Repayment releases the cNFT and returns the asset. Default triggers a 14-day Dutch auction; proceeds settle the loan."),
]
for i, (num, title, desc) in enumerate(steps):
    cell = steps_tbl.rows[0].cells[i]
    set_cell_bg(cell, "F8F6F2")
    set_cell_borders(cell, top="single", bottom="nil", left="nil", right="nil",
                      color="0E7C7B", size=18)
    set_cell_padding(cell, top=70, bottom=70, left=100, right=100)
    # number + title row
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run(f"{num}  ")
    r.font.name = "Courier New"
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUTE
    r.font.bold = True
    r = p1.add_run(title)
    r.font.name = "Helvetica"
    r.font.size = Pt(8.5)
    r.font.color.rgb = INK_DARK
    r.font.bold = True
    # description
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(desc)
    r.font.name = "Helvetica"
    r.font.size = Pt(7.5)
    r.font.color.rgb = INK


# ============================================================================
# 4. UNIT ECONOMICS  (2 columns: Borrower | Capital Providers)
# ============================================================================
section_header(doc, 4, "Unit economics")

econ_tbl = doc.add_table(rows=1, cols=2)
econ_tbl.autofit = False

# helper — fill an econ cell
def fill_econ_col(cell, header, rows):
    set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")
    set_cell_padding(cell, top=20, bottom=20, left=120, right=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(header)
    r.font.name = "Helvetica"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = INK_DARK
    # Row table inside cell for label/value alignment
    rt = cell.add_table(rows=len(rows), cols=2)
    rt.autofit = False
    for ridx, (label, val) in enumerate(rows):
        lcell = rt.rows[ridx].cells[0]
        vcell = rt.rows[ridx].cells[1]
        for cc in (lcell, vcell):
            set_cell_borders(cc, top="nil", bottom="nil", left="nil", right="nil")
            set_cell_padding(cc, top=20, bottom=20, left=0, right=0)
        lp = lcell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        r = lp.add_run(label)
        r.font.name = "Helvetica"; r.font.size = Pt(8.5); r.font.color.rgb = INK
        vp = vcell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vp.paragraph_format.space_after = Pt(0)
        r = vp.add_run(val)
        r.font.name = "Helvetica"; r.font.size = Pt(8.5); r.font.bold = True
        r.font.color.rgb = ACCENT

borrower_rows = [
    ("All-in APR (2% / month)",       "24%"),
    ("LTV (vs Caixa 20% scrap)",      "Up to 50%"),
    ("Term",                          "3-month, renewable"),
    ("vs credit-card revolving",      "16× cheaper"),
]
capital_rows = [
    ("Senior tranche (USDC, fixed)",       "8% APR"),
    ("Junior tranche (USDC, fixed)",       "12% APR"),
    ("Vaulx protocol-owned first-loss",    "5% buffer"),
    ("vs Maple syrupUSDC senior",          "+100 bps"),
]

fill_econ_col(econ_tbl.rows[0].cells[0], "BORROWER", borrower_rows)
fill_econ_col(econ_tbl.rows[0].cells[1], "CAPITAL PROVIDERS", capital_rows)

# Vaulx revenue band (single full-width row table)
rev_tbl = doc.add_table(rows=1, cols=1)
rcell = rev_tbl.rows[0].cells[0]
set_cell_bg(rcell, "EFEDE5")
set_cell_borders(rcell, top="nil", bottom="nil", left="single", right="nil",
                  color="0E7C7B", size=18)
set_cell_padding(rcell, top=80, bottom=80, left=140, right=140)
p = rcell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
r = p.add_run("VAULX REVENUE   ")
r.font.name = "Helvetica"; r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = ACCENT
r = p.add_run("~$300–600 per asset / year (6–12% of borrowed). Origination + servicing + curator fees.")
r.font.name = "Helvetica"; r.font.size = Pt(8.5); r.font.color.rgb = INK_DARK


# ============================================================================
# 5. WHERE MERCADO BITCOIN FITS
# ============================================================================
section_header(doc, 5, "Where Mercado Bitcoin fits")

add_body(doc,
    "Vaulx has built the protocol, the custody layer, and the borrower funnel. To operate compliantly "
    "in Brazil at scale, we need a regulated counterparty. Three areas where MB and Vaulx could cooperate:",
    size=9, space_before=0, space_after=4,
)

mb_areas = [
    ("Regulated tokenization",
     "MB issues the on-chain representation (cNFT) of the vaulted asset under its existing license."),
    ("Fiat rails · BRL ↔ USDC",
     "MB acts as the primary on/off-ramp for borrower disbursement and repayment via Pix."),
    ("Yield product distribution",
     "Vaulx senior tranche (8% USDC, physical collateral, insured) listed for MB institutional / retail."),
]
for label, desc in mb_areas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Mm(2)
    # square bullet (text)
    r = p.add_run("▪  ")
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.color.rgb = ACCENT; r.font.bold = True
    # label
    r = p.add_run(label + "  ")
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = INK_DARK
    # description
    r = p.add_run(desc)
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.color.rgb = INK


# ============================================================================
# 6. WHAT PROTECTS EVERYONE  (safeguards — no commercial terms)
# ============================================================================
section_header(doc, 6, "What protects everyone")

safeguards = [
    ("Atomic invariant",
     "No USDC disburses until the licensed custodian signs custody-confirmation, atomically, in the same transaction."),
    ("Independent licensed custody",
     "Brinks- / Loomis-class custodian — not Vaulx, not MB — retains physical control of the asset."),
    ("Lloyd's-class insurance",
     "Theft + damage to the trustee on every vaulted asset."),
    ("5% protocol-owned first-loss",
     "Vaulx posts capital below LP tranches — we eat default losses first."),
]
for label, desc in safeguards:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Mm(2)
    r = p.add_run("▪  ")
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.color.rgb = ACCENT; r.font.bold = True
    r = p.add_run(label + "  ")
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = INK_DARK
    r = p.add_run(desc)
    r.font.name = "Helvetica"; r.font.size = Pt(9); r.font.color.rgb = INK


# ============================================================================
# FOOTER  (team line + URLs)
# ============================================================================
# Hairline before footer
foot_rule = doc.add_paragraph()
foot_rule.paragraph_format.space_before = Pt(8)
foot_rule.paragraph_format.space_after = Pt(2)
add_horizontal_rule(foot_rule, color="D8D8DB", size=4)

team_p = doc.add_paragraph()
team_p.paragraph_format.space_before = Pt(0)
team_p.paragraph_format.space_after = Pt(2)
r = team_p.add_run("TEAM:  ")
r.font.name = "Helvetica"; r.font.size = Pt(7); r.font.bold = True; r.font.color.rgb = INK_DARK
r = team_p.add_run("George Dimitrov  CEO/CTO   ·   Marcelo  COO  (Gitel — 38 yrs BR electronic-security)   ·   "
                    "Rodrigo  Head of BD   ·   Edson  Sr. Solana Eng.   ·   Felipe  DeFi Advisor (4p.finance)")
r.font.name = "Helvetica"; r.font.size = Pt(7); r.font.color.rgb = MUTE

# Footer URL line — 2-col table for left/right alignment
foot_tbl = doc.add_table(rows=1, cols=2)
foot_tbl.autofit = False
foot_tbl.columns[0].width = Mm(90)
foot_tbl.columns[1].width = Mm(90)

lcell = foot_tbl.rows[0].cells[0]
set_cell_borders(lcell, top="nil", bottom="nil", left="nil", right="nil")
set_cell_padding(lcell, top=0, bottom=0, left=0, right=0)
lp = lcell.paragraphs[0]
lp.paragraph_format.space_before = Pt(0); lp.paragraph_format.space_after = Pt(0)
r = lp.add_run("vaulx.fi   ·   github.com/Vaulxfi")
r.font.name = "Helvetica"; r.font.size = Pt(7.5); r.font.bold = True; r.font.color.rgb = INK_DARK

rcell = foot_tbl.rows[0].cells[1]
set_cell_borders(rcell, top="nil", bottom="nil", left="nil", right="nil")
set_cell_padding(rcell, top=0, bottom=0, left=0, right=0)
rp = rcell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_before = Pt(0); rp.paragraph_format.space_after = Pt(0)
r = rp.add_run("Built. Tested. Live on Solana Devnet today.   ·   STRICTLY CONFIDENTIAL")
r.font.name = "Helvetica"; r.font.size = Pt(7); r.font.color.rgb = MUTE


# === Save ================================================================
output = "Vaulx_MB_OnePager.docx"
doc.save(output)
print(f"Generated: {output}")
